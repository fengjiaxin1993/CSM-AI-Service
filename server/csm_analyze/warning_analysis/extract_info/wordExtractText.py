import json
from docx import Document
from server.csm_analyze.warning_analysis.extract_info.helper import clean_text


class WORDExtractText:
    """电力行业告警报告Word文档解析器"""

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.tables_data = self.__extract_tables()
        self.full_text = self.__read_docx()

    def __extract_tables(self) -> list:
        """提取表格数据，保留列名和值的对应关系"""
        tables_data = []
        for table in self.doc.tables:
            table_data = []
            headers = []
            for i, row in enumerate(table.rows):
                row_cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    # 第一行作为表头
                    headers = row_cells
                else:
                    # 数据行，与表头对应
                    row_dict = {}
                    for j, cell_text in enumerate(row_cells):
                        if j < len(headers):
                            row_dict[headers[j]] = cell_text
                        else:
                            row_dict[f"列{j+1}"] = cell_text
                    if row_dict:
                        table_data.append(row_dict)
            if table_data:
                tables_data.append(table_data)
        return tables_data

    def __read_docx(self) -> str:
        """读取Word文档段落文本（不含表格）"""
        full_text = ""
        # 只读取段落，不读取表格
        for para in self.doc.paragraphs:
            full_text += para.text + "\n"
        full_text = clean_text(full_text)
        return full_text.strip()


# ---------------------- 测试代码 ----------------------
if __name__ == "__main__":
    # 测试Word解析
    DOCX_PATH1 = "../test_file/关于110kVXX变告警说明.docx"
    DOCX_PATH2 = "../test_file/附件1：关于A二级水电站2月3日的网络安全事件调查报告5 - v2【违规外联】.docx"
    DOCX_PATH3 = "../test_file/附件1：说明模板-关于XX变03月18日的告警情况说明-V2-WLAQ2026031801.docx"
    DOCX_PATH4 = "../test_file/附件1：说明模板-关于XX风光电厂）05月22日的告警情况说明5.31修改版本.docx"
    parser = WORDExtractText(DOCX_PATH2)
    print("=== 段落文本 ===")
    print(parser.full_text)
    print("\n=== 表格数据 ===")
    for i, table in enumerate(parser.tables_data):
        print(f"表格 {i+1}:")
        print(json.dumps(table, ensure_ascii=False, indent=2))
