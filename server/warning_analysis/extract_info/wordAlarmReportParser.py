# -*- coding: utf-8 -*-
import json
from typing import Dict, Optional
import re
from docx import Document
from server.warning_analysis.extract_info.helper import _init_structured_fields, clean_text


class WORDAlarmReportParser:
    """电力行业告警报告Word文档解析器"""

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.table_dict = self.__extract_table_dic()
        self.full_text = self.__read_docx()
        self.structured_data = _init_structured_fields()

    def __read_docx(self) -> str:
        """读取Word文档全文"""
        full_text = ""
        # 读取所有段落，传入 doc.part 以获取编号
        for para in self.doc.paragraphs:
            # 获取编号
            full_text += para.text + "\n"
        return clean_text(full_text).strip()

    def __extract_table_dic(self) -> dict:
        """提取Word文档中的表格"""
        for table in self.doc.tables:
            table_dict = {}
            # 假设第一行为表头
            headers = [clean_text(cell.text) for cell in table.rows[0].cells]
            # 读取数据行
            row = table.rows[1]
            for idx, cell in enumerate(row.cells):
                if idx < len(headers):
                    table_dict[headers[idx]] = clean_text(cell.text)
            return table_dict

    def __extract_section_content(self, section_name: str, next_section_name: Optional[str] = None) -> str:
        """提取指定模块的内容"""
        if not section_name and not next_section_name:
            return ""
        elif not section_name and next_section_name:
            pattern = rf"\s*(.+?)\s*{next_section_name}"
        elif section_name and not next_section_name:
            pattern = rf"{section_name}\s*(.+?)$"
        else:
            pattern = rf"{section_name}\s*(.+?)\s*{next_section_name}"

        match = re.search(pattern, self.full_text, re.DOTALL)
        if match:
            content = match.group(1).strip()
            content = re.sub(r"\s+", " ", content)
            return content
        return ""

    def parse(self) -> Dict:
        """执行完整解析流程"""
        # 1. 提取标题
        self.structured_data["报告标题"] = self.__extract_section_content("", "告警信息")
        # 2. 提取告警信息
        self.structured_data["告警信息"] = self.__extract_section_content("告警信息", "处置过程")
        # 提取告警表格信息
        if self.table_dict:
            self.structured_data["设备名称"] = self.table_dict.get("设备名称", "")
            self.structured_data["设备类型"] = self.table_dict.get("设备类型", "")
            self.structured_data["告警时间"] = self.table_dict.get("告警时间", "")
            self.structured_data["告警内容"] = self.table_dict.get("告警内容", "")
        # 3. 提取处置过程
        self.structured_data["处置过程"] = self.__extract_section_content("处置过程", "原因分析")
        # 4. 提取原因分析
        self.structured_data["原因分析"] = self.__extract_section_content("原因分析", "整改情况")
        # 5. 提取整改情况
        self.structured_data["整改情况"] = self.__extract_section_content("整改情况", "防范措施")
        # 6. 提取防范措施
        self.structured_data["防范措施"] = self.__extract_section_content("防范措施")
        return self.structured_data


# ---------------------- 测试代码 ----------------------
if __name__ == "__main__":
    # 测试PDF解析（普通PDF）
    PDF_PATH1 = "../test_file/关于110kVXX变告警说明.docx"
    PDF_PATH2 = "../test_file/附件1：关于A二级水电站2月3日的网络安全事件调查报告5 - v2【违规外联】.docx"
    PDF_PATH3 = "../test_file/附件1：说明模板-关于XX变03月18日的告警情况说明-V2-WLAQ2026031801.docx"
    PDF_PATH4 = "../test_file/附件1：说明模板-关于XX风光电厂）05月22日的告警情况说明5.31修改版本.docx"
    parser = WORDAlarmReportParser(PDF_PATH1)
    result = parser.parse()
    print("=== Word文档解析结果 ===")
    print(json.dumps(result, ensure_ascii=False, indent=2))

# === Word文档解析结果 ===
# {
#   "报告标题": "·XX地调关于X月X日XX变的告警情况说明",
#   "告警信息": "该告警被XX一平面II区监测装置拦截并上报至地调网络安全监视平台，后上传至省调网络安全监测平台，告警内容为由综自后台1（XXXX）拦截XXXX的XXX端口向XXX的X端口之间存在XX端口访问。",
#   "设备名称": "XXX监测装置",
#   "设备类型": "监测装置",
#   "告警时间": "2025-X-XX:X:X",
#   "告警内容": "由综自后台1（XX）拦截XXXX的XX端口向XXX的XX端口之间存在XX端口访问",
#   "处置过程": "1、在收到省调安防的电力监控系统网络安全隐患整改通知单后，XX地调自动化于x月x日上午联系XX变所属检修单位，告知XX期间无检修计划安排且检修人员安规未考试，计划X日进行现场处置。2、X日上午，检修人员到达现场通过查看综自后台1运行情况发现7号该台电脑存在自动重启现象，且现场后台主机所运行的南瑞继保监控程序至14日都未正常运行。如下图；图1：设备运行情况3、在确定该台主机存在重启后且现场主机业务系统厂家南瑞继保所采用SQLServer2000使用NetBios协议进行通讯，业务系统启动时会自动启用轮询，监控客户端和服务器端的数据收发区域，测试读取和写入功能得过程，通过查看该台主机网卡配置，发现该主机同时存在A/B网段IP（xxx），且均处于在用状态，因此会对该网段全部主机进行随机端口扫描。同时查看该主机探针网络白名单，发现xx不在该白名单内，如图2所示，随即自动化值班人员通过网安平台添加白名单后（图3）告警消除。图2：白名单未添加图3：添加白名单",
#   "原因分析": "(1)经变电检修人员现场检查，根据后台主机监控系统运行时常，判断出杨家湾变电站综自后台主机1于7号出现过自动重启现象，该台主机系统为WindowsXP系统，存在系统老旧、内存较低、长久未更新等因素，长时间运行后存在程序无响应、死机等状况，因此业务系统厂家南瑞继保根据主机监控程序运行情况设置了自动重启机制。(2)该站综自后台业务系统均为南瑞继保公司系统，商用库采用XXXX，由于XXXXX使用XXXX协议进行通讯，因此业务系统启动时会自动启用轮询，监控客户端和服务器端的数据收发区域，测试读取和写入功能（针对该站后台主机所在A/B网段XXXXX进行全网扫描），造成多次由综自后台主机1采用随机端口（XXX）访问综自后台主机2（XXXX）目的端口XXX端口的告警产生。（3）本次告警所访问的网段XXXX在宗自后台主机1中未添加，造成白名单缺失，从而在设备自动轮巡时造成告警产生，包括采用随机端口XX轮询的告警，被检测装置拦截判断为紧急告警上报至网安平台。",
#   "整改情况": "已通过地调网安平台添加后台综自主机1B网IP地址白名单。如下图：",
#   "防范措施": "1、强化运维管理。组织学习电力监控系统网络安全相关管理规定，修编完善设备巡视、现场作业指导卡等。2、针对老旧存在漏洞等操作系统进行国产化替代，通过改造，使用安全性更好的国产化操作系统和自主可控硬件平台。"
# }
